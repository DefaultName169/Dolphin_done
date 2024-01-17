import zipfile
import argparse
import re
import os
import docx
import shutil
# import docx.oxml.text.CT_P
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

parser = argparse.ArgumentParser()
parser.add_argument('--wordfile', '--w', required=False, help="location of word file")
parser.add_argument('--pathfolder', '--dir', required=False, help="")
parser.add_argument('--pathfile', '--file', required=False, help="")
# parser.add_argument('--name', default=None, required=False)
cli_args = parser.parse_args()

def get_input(var_name):
    auto_input = getattr(cli_args, var_name, None) 
    if auto_input :
        print("Auto input:", auto_input)
        return auto_input
    else:
        return input("Manual input: ")

woldfile = get_input("wordfile")
pathfolder = get_input("pathfolder")
pathfile = get_input("pathfile")

option_change = open('/data/projects/memory_compiler2/REPOSITORY/Rampiler_doc/database/list_change', 'r').readlines()
option_change = [i.strip() for i in option_change]

print(option_change)

backfile = '/data/projects/memory_compiler2/scripts/fire_stim_script/makeDoc/back.docx'

def type_xml(xml_index):
    if re.search(r'<w:pStyle w:val="Caption"/>.*SEQ', xml_index):
        return 'caption'

    if re.search(r'<w:p\s.*<v:imagedata|<w:p>.*<v:imagedata|<w:drawing>.*<w:t\s|<w:drawing>.*<w:t>' , xml_index):
        return 'notchange'    
    
    drawing = re.findall(r'<w:drawing>.*?</w:drawing>',xml_index)
    for draw in drawing:
        if not re.search(r'"rId\d+"',draw):
            return 'notchange'
    
    if re.search(r'<w:drawing>' , xml_index) :
        return 'img'
    
    if re.search(r'<w:pPr><w:pStyle w:val=\"Heading\d\"/>.*?</w:p>', xml_index):
        return 'heading'
    
    if re.search(r'<secPr',xml_index):
        return 'last'
    
    this_type = re.sub(r'(^<w:)(p|tbl)(.*)', r'\2', xml_index)
    return this_type

def readdocx(path):
    dirtrx = './removedup'
    doc = docx.Document(path)
    if os.path.exists(pathfolder) :
        shutil.rmtree(pathfolder)

    os.makedirs(pathfolder)

    with zipfile.ZipFile(path, 'r') as zipdoc:
        for file in zipdoc.namelist():
            if re.search('word/media',file):
                zipdoc.extract(file,path=dirtrx)
        data_xml_doc = zipdoc.read('word/document.xml').decode('utf8')
        string_link = zipdoc.read('word/_rels/document.xml.rels').decode('utf8')

    list_img = re.findall(r'<Relationship Id="(\w+)" Type=.*? Target="(.*?)"/>',string_link)
    content_type_full = {'emf':'<Default Extension="emf" ContentType="image/x-emf"/>',
                        'jpeg':'<Default Extension="jpeg" ContentType="image/jpeg"/>',
                        "png" : '<Default Extension="png" ContentType="image/png"/>',
                        "rels":'<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
                        "vsd":'<Default Extension="vsd" ContentType="application/vnd.visio"/>',
                        "vsdx":'<Default Extension="vsdx" ContentType="application/vnd.ms-visio.drawing"/>',
                        "xml":'<Default Extension="xml" ContentType="application/xml"/>'}
    skip_heading = ['LIST OF TABLES','LIST OF FIGURES', 'TABLE OF CONTENTS']
    path_print = '__path__ = '+ woldfile +'\n\n\nTREE:\n'
    start = False
    last_type = ''
    path_save = ''
    path_of_image = ''
    path_to_mkdir = ''
    doc_name = 0
    arr_level = [0]
    heading_notcount = 0
    new_doc = None
    add_caption_to_next = False
    skip = False
    run = False
    lvlrun = 0
    
    for p in doc.paragraphs:
        for run in p.runs:
            run = run._r
            if '<w:br w:type="page"/>' in run.xml:
                run.getparent().remove(run)
                
    for doc_inside in doc.element.body:
        if str(type(doc_inside)) == '''<class 'lxml.etree._Element'>''':
            continue
        xml = re.sub(r'xmlns:\w+=".*?"\s|\n\s+|\n','',doc_inside.xml)
        xml = re.sub(r'\sxmlns:\w+=".*?">','>', xml)

        if re.search(r'<w:pStyle w:val="Heading\d"/>', xml):
            rs = doc_inside.xpath('.//w:t')
            str_heading = u"".join([r.text for r in rs])
            level_heading = re.sub(r'(.*w:val="Heading)(\d)(".*)', r'\2' , xml)
            if not any(element in str_heading for element in skip_heading):
                start = True

        if start:
            this_type = type_xml(xml)
            if not re.search(r'<w:t>|<w:t\s|"rId\d+"',xml):
                # if not last_type == 'p':
                #     last_type = ''
                continue
###############################################################################

            if last_type == 'p' and this_type != 'p':
                new_doc.element.body.remove(new_doc.element.body[0])
                
                if re.search(r'<w:p>|<w:p\s|<w:tbl>"', new_doc.element.body.xml) :
                    doc_name += 1
                    path_save = pathfolder + path_to_mkdir + '/' + str(doc_name) + '.docx'
                    new_doc.save(path_save)
                
                    if re.search(r'^\t', path_in_txt_str) :
                        path_print += re.sub(r'(^\t+)(.*)', r'\1', path_in_txt_str) + '\t' + path_save.split('/')[-1] +'\n'
                    else :
                        path_print += '\t'+ path_save.split('/')[-1] +'\n'
                       
                        
#################################################################################################
            if this_type == 'heading':
                doc_name = 0
                if last_type == '':
                    new_doc = docx.Document(backfile)
                path_to_mkdir = ''

                if int(level_heading) == len(arr_level) : 
                    arr_level[-1] += 1
                elif int(level_heading) > len(arr_level):
                    arr_level.append(1)
                else:
                    arr_level = arr_level[:int(level_heading)]
                    arr_level[-1] += 1

                path_in_txt_str = ''
                
                arr_path = [i for i in arr_level]
                arr_path[0] -= heading_notcount
                path_to_mkdir = '/' + '/'.join(map(str,arr_level)) + '/'

                if re.search(r'Appendix \w+:|REVISION HISTORY', str_heading):
                    heading_notcount += 1
                    path_in_txt_str = '\t'*(len(arr_path) - 1) + str(arr_level[-1]) + ' : ' + str_heading 
                else :
                    path_in_txt_str = '\t'*(len(arr_path) - 1) + str(arr_level[-1]) + ' : ' + '.'.join(map(str,arr_path)) + '. '+ str_heading 

                if level_heading == lvlrun:
                    run = False 

                if not re.search(r'tsmc_total', pathfolder):
                    if str_heading in option_change:
                        run = True
                        lvlrun = level_heading
                        last_type = this_type
                else :
                    run = True

                if run :
                    print(path_in_txt_str)
                    path_print += path_in_txt_str + '\n'
                last_type = ''
                new_doc = docx.Document(backfile)
                continue
###############################################################################
            if not run: 
                continue
###############################################################################
            if not os.path.exists(pathfolder + path_to_mkdir):
                os.makedirs(pathfolder + path_to_mkdir)


###############################################################################

            if this_type == 'tbl':
                doc_name += 1
                path_save = pathfolder + path_to_mkdir + '/' + str(doc_name) + '.docx'
                new_doc = docx.Document(backfile)
                new_doc.element.body[-2].addnext(doc_inside)
                new_doc.element.body.remove(new_doc.element.body[0])
                new_doc.save(path_save)
                
                if re.search(r'^\t', path_in_txt_str) :
                    path_print += re.sub(r'(^\t+)(.*)', r'\1', path_in_txt_str) + '\t' + path_save.split('/')[-1] +'\n'
                else :
                    path_print += '\t'+ path_save.split('/')[-1] +'\n'


###############################################################################

            if this_type == 'caption':
                rs = doc_inside.xpath('.//w:t')
                has_caption = u"".join([r.text for r in rs])
                type_caption = re.findall(r'(Table|Figure)\s+\d+\s?\W?\s?', has_caption)[0]
                if(type_caption == 'Table'):
                    add_caption_to_next = True
                    skip = True

                if(type_caption == 'Figure'):
                    path_print = path_print[:-1]
                    path_print += " : " + re.sub(r'Figure(\s+)(\d+)(\s?[\.:]\s?)','Figure'+ r'\1' + r'\\f' + r'\3' , has_caption) + '\n'
            
                if has_caption and re.search(r'^Appendix\s?\w', has_caption):
                    has_caption = None

                   
###############################################################################

            if this_type == 'img':
                doc_name += 1
                id_img = re.sub(r'(.*<a:blip r:embed=")(.\w+)(.*)',r'\2', xml)
                
                for img in list_img:
                    if img[0] == id_img :
                        name_pic = re.sub(r'(.*)\d+', str(doc_name), img[1])
                        path_save = pathfolder + path_to_mkdir + '/' + name_pic
                        path_of_image = 'word/'+ img[1]
                        last_type = this_type
                        break
                
                if re.search(r'^\t', path_in_txt_str) :
                    path_print += re.sub(r'(^\t+)(.*)', r'\1', path_in_txt_str) + '\t' + path_save.split('/')[-1] +'\n'
                else :
                    path_print += '\t'+ path_save.split('/')[-1] +'\n'

                shutil.copy(dirtrx + '/'+ path_of_image, path_save)

                    
###############################################################################

            if this_type == 'p':
                if last_type != 'p':
                    new_doc = docx.Document(backfile)
                new_doc.element.body[-2].addnext(doc_inside)

###############################################################################

            if this_type == 'notchange':
                doc_name += 1
                new_doc = docx.Document(backfile)
                path_save = pathfolder + path_to_mkdir + '/' + str(doc_name) + '.docx'

                new_doc.element.body[-2].addnext(doc_inside)
                new_doc.save(path_save)
                
                with zipfile.ZipFile(path_save, 'a') as zipoutput:
                    data_xml_new_doc = zipoutput.read('word/document.xml').decode('utf-8')
                    data_xml_rels_new_doc = zipoutput.read('word/_rels/document.xml.rels').decode('utf-8')
                    content_type = zipoutput.read('[Content_Types].xml').decode('utf-8')
                    
                    rid_need_change =  re.findall(r'="rId(\d+)"', xml)
                    rid_need_change = list(dict.fromkeys(rid_need_change))
                    list_img_id = re.findall(r'(<Relationship Id="rId)(.*?)(".*?Target=")(.*?)("\/>)', data_xml_rels_new_doc)
                    num_id_img_exist = []
                    for i in list_img_id:
                        num_id_img_exist.append(int(i[1]))
                    num_id_img_exist.sort()

                    num_last_id = num_id_img_exist[-1]
                    num_last_id_path = 2

                    for rid in rid_need_change :
                        num_last_id += 1
                        num_last_id_path += 1
                        data_xml_new_doc = re.sub(r'="rId'+ rid + '"', r'="rId'+ str(num_last_id) + '"', data_xml_new_doc)

                        data_xml_rels = re.findall(r'<Relationship Id="rId'+ rid + '"' + r'.*?Target=".*?"/>',  string_link)[0]
                        
                        target = re.search(r'<Relationship Id="rId'+ rid+ '"' + r'.*?Target="(.*?)\.(.*?"/>)', data_xml_rels).group(1)
                        # print(target)
                        if not re.search(r'\d+', target):
                            data_xml_rels_to_doc = re.sub(r'<Relationship Id="rId'+ rid+ '"' + r'(.*?)Target="(.*?)(\..*?"/>)' , r'<Relationship Id="rId'+ str(num_last_id)+ '"' + r'\1Target="\2 ' + str(num_last_id_path) + r'\3',  data_xml_rels)
                        else:
                            data_xml_rels_to_doc = re.sub(r'<Relationship Id="rId'+ rid+ '"' + r'(.*?)Target="(.*?)\d+(\..*?"/>)' , r'<Relationship Id="rId'+ str(num_last_id)+ '"' + r'\1Target="\2 ' + str(num_last_id_path) + r'\3',  data_xml_rels)
                        data_xml_rels_to_doc = re.sub(r'Target="(.*)\s(.*)', r'Target="\1\2', data_xml_rels_to_doc)
                        data_xml_rels_new_doc = re.sub(r'</Relationships>', data_xml_rels_to_doc + '</Relationships>', data_xml_rels_new_doc)
                        # print(data_xml_rels_new_doc)
                        
                        path_of_img = re.findall(r'<Relationship Id="rId'+ rid + r'.*?Target="(.*?)"/>',  string_link)[0]
                        # print(path_of_img)

                        if not re.search(r'\d+', target):
                            new_path_img = re.sub(r'\.', str(num_last_id_path) + '.' , path_of_img)
                        else:
                            new_path_img = re.sub(r'\d+\.', str(num_last_id_path) + '.' , path_of_img)
                        # print(new_path_img)
                        type_content = re.sub(r'(.*)\.(\w+)',r'\2', path_of_img)

                        if not re.search(content_type_full[type_content],content_type):
                            content_type = re.sub(r'<Types(.*?)>',r'<Types\1>'+ content_type_full[type_content] ,content_type)

                        zipoutput.writestr('word/' + new_path_img, zipfile.ZipFile(path, 'r').open('word/' + path_of_img).read())

                    zipoutput.writestr('word/document.xml', data_xml_new_doc)
                    zipoutput.writestr('word/_rels/document.xml.rels', data_xml_rels_new_doc)
                    zipoutput.writestr('[Content_Types].xml', content_type)

                new_doc = docx.Document(path_save)
                new_doc.element.body.remove(new_doc.element.body[0])
                new_doc.save(path_save)   
                
                if re.search(r'^\t', path_in_txt_str) :
                    path_print += re.sub(r'(^\t+)(.*)', r'\1', path_in_txt_str) + '\t' + path_save.split('/')[-1] +'\n'
                else :
                    path_print += '\t'+ path_save.split('/')[-1] +'\n'

####################################################################################

            if add_caption_to_next:
                if skip:
                    skip = False
                else:
                    path_print = path_print[:-1]
                    path_print += " : " + re.sub(r'Table(\s+)(\d+)(\s?[\.:]\s?)','Table'+ r'\1' + r'\\t' + r'\3' , has_caption) + '\n'
                    add_caption_to_next = False

            last_type = this_type
######################################################################################
    #if os.path.exists('pathname.txt') : 
    #    os.remove('pathname.txt')
    path_to_print = open(pathfile , 'w')
    path_to_print.write(path_print)
    print(path_print)
    path_to_print.close()
    shutil.rmtree(dirtrx)

readdocx(woldfile)

#def xmlstring_to_oxmlelement(string):
#    queue = []
#    save = []
#    eval = ''
#    first = ()
#    array = re.finditer('<.+?>', string)
#    index_last_text = 0
#    for x in array:
#        if x == '':
#            continue
#        else :
#            text = x.group()
#            if queue:         
#                first = re.sub(r'.*:(.*)', r'\1', queue[0])    
#                if queue[-1] == 'w:t':
#                    eval += 't.text = \'\'\'' + string[index_last_text:x.start()] + '\'\'\'' + '\n'
#
#                if re.search('/' + queue[-1], text) :
#                    a = re.sub(r'.*:(.*)', r'\1', queue.pop())
#                    while save:
#                        eval += a + '.append(' + re.sub(r'.*:(.*)', r'\1', save.pop(0)) + ')' + '\n'
#                    if queue:
#                        eval += re.sub(r'.*:(.*)', r'\1', queue[-1]) + '.append(' + a + ')' + '\n'
#                    continue
#            
#            text = re.sub('<|>', '', text)
#            pins = text.split(' ',1)
#            name_clr = re.sub('/', '', pins[0])
#            name = re.sub(r'.*:(.*)', r'\1', name_clr)
#            eval += name + ' = OxmlElement(\'' + name_clr + '\')'+'\n'
#            queue.append(pins[0])
#
#            if len(pins) > 1:
#                matches = re.findall(r'\s?(.*?)=\"(.*?)\"', pins[1])
#                for match in matches :
#                    key = match[0]
#                    value = match[1]
#                    eval += name + '.set(qn(\'' + key + '\'),\'' + value + '\')' +'\n'
#                
#            if pins[0] == 'w:t':
#                index_last_text = x.end()
#
#            if re.search(r'\/', pins[-1]) :
#                pin = re.sub('/', '' , pins[0])
#                save.append(pin)
#                queue.pop()
#            
#    eval = 'output = ()\n' + eval + 'output = ' + first
#    loc = {}
#    exec(eval, globals(), loc)
#    output = loc['output']
#    return output
#
#def type_xml(xml_index):
#    if re.search(r'<w:pStyle w:val="Caption"/>.*SEQ', xml_index):
#        return 'caption'
#    if re.search(r'<w:drawing>|<w:p .*<v:imagedata|<w:p>.*<v:imagedata' , xml_index) :
#        return 'img'
#    # if xml_index == '' :
#    #     return 'p'
#    this_type = re.sub(r'(^<w:)(p|tbl)(.*)', r'\2', xml_index)
#    return this_type
#
#
#
#def readdocx(path):
#    dirtrx = 'removedup'
#    with zipfile.ZipFile(path, 'r') as zip:
#        zip.extractall(path=dirtrx)
#        string = zip.read('word/document.xml').decode('utf8')
#        string_link = zip.read('word/_rels/document.xml.rels').decode('utf8')
#
#    list_img = re.findall(r'<Relationship Id="(\w+)" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/(\w+\.\w+)"/>',string_link)
#
#    xml_heading = re.finditer( r'(<w:pPr><w:pStyle w:val=\"Heading\d\"/>.*?</w:p>)', string) 
#    xml_heading = itertools.chain(xml_heading, '$')
#    
#    
#    last_xml_index = 0
#
#    skip_heading = ['REVISION HISTORY','LIST OF TABLES','LIST OF FIGURES', 'TABLE OF CONTENTS']
#    str_heading = ''
#    xml_inside = ''
#    path_print = '__path__ = '+ pathfolder +'\nname_of_output = output\n\n\nTREE:\n'
#    arr_level = [0]    
#    i = 0
#
#
#    for xml in xml_heading:
#        if any(element in str_heading for element in skip_heading):
#            str_heading = ''
#            continue
#
#        if xml == '$' :
#            xml_inside = string[last_xml_index : len(string)] 
#        else:
#            xml_inside = string[last_xml_index : xml.span()[0]] 
#
#        if str_heading != '':
#            xml_inside = re.sub(r'<w:lastRenderedPageBreak/>|<w:proofErr w:type="spellStart"/>|<w:proofErr w:type="spellEnd"/>|<w:bookmarkStart.*?/>|<w:proofErr.*?/>', '', xml_inside)
#            all_xml_inside = re.findall(r'(<w:p>.*?</w:p>|<w:p\s.*?</w:p>|<w:tbl>.*?</w:tbl>)',xml_inside)
#
#            if int(level_heading) == len(arr_level) : 
#                arr_level[-1] += 1
#            elif int(level_heading) > len(arr_level):
#                arr_level.append(1)
#            else:
#                arr_level = arr_level[:int(level_heading)]
#                arr_level[-1] += 1
#
#            path_in_txt_str = ''
#            path_to_mkdir = ''
#
#            for i in range(0 , len(arr_level)):
#                if i == len(arr_level) - 1:
#                    path_in_txt_str += str(arr_level[i]) +' : ' + str_heading
#                    path_to_mkdir += '/' + str(arr_level[i])
#                    break
#                path_in_txt_str += '\t'
#                path_to_mkdir += '/' + str(arr_level[i]) 
#
#            print(path_in_txt_str)
#
#
#            if not re.search(r' Appendix \w+:',path_in_txt_str):
#                numbering = re.sub('^/','', path_to_mkdir)
#                numbering = numbering.replace('/','.')
#                path_in_txt_str = re.sub(r'(\d : )(.*)', r"\1 "+ numbering + r". \2",path_in_txt_str)
#                
#            path_print += path_in_txt_str + '\n'
#
#
#            if os.path.exists(pathfolder + path_to_mkdir) : 
#                shutil.rmtree(pathfolder + path_to_mkdir)
#            os.makedirs(pathfolder + path_to_mkdir)
#            
#            last_type = 'new'
#            doc_name = 0
#            doc = None
#            has_caption = None
#            path_save = ''
#            path_of_image = ''
#            caption = []
#            is_tabcap = False
#
#            for x in range(0 , len(all_xml_inside) + 1) :
#                if x < len(all_xml_inside) :
#                    xml_index = all_xml_inside[x] 
#                    if re.search(r'^<w:p',xml_index) : 
#                        xml_i = re.finditer(r'<w:p |<w:p>',xml_index)
#                        *_, last = xml_i 
#                        xml_index = str(xml_index)[last.span()[0] : len(xml_index)]
#                else :
#                    xml_index = ''
#
#                this_type = type_xml(xml_index)
#
#                if not( re.search(r'<w:t>|<w:t ',xml_index)) and this_type == 'p' :
#                    continue
#                
#                if this_type == 'caption':
#                    has_caption = re.findall(r'<w:t.*?>(.*?)</?w', xml_index)
#                    has_caption = ''.join([str(elem) for elem in has_caption])
#                    has_caption = re.sub(r'Table\s+\d+\s?\W?\s?|Figure\s+\d+\s?\W?\s?','', has_caption)
#                    caption.append(has_caption)
#                    if has_caption and re.search(r'^Appendix \w', has_caption):
#                        has_caption = None
#                    
#                    # if last_type == 'tbl' and caption:
#                    #     print('_____________________________________')
#                    #     path_print = path_print[:-1]
#                    #     path_print += " : " + caption.pop(0) + '\n' 
#                    #     has_caption = None
#                    continue
#                
#                if this_type != last_type or this_type == 'img':
#                    if(path_save != ''):
#                        if re.search(r'^\t', path_in_txt_str) :
#                            path_print += re.sub(r'(^\t+)(.*)',r'\1',path_in_txt_str) + '\t' + path_save.split('/')[-1] +'\n'
#                        else :
#                            path_print += '\t'+ path_save.split('/')[-1] +'\n'
#                    
#                    if is_tabcap :
#                        print('-----------------------------------------')
#                        path_print = path_print[:-1]
#                        path_print += " : " + caption.pop(0) + '\n' 
#                        has_caption = None
#                        is_tabcap = False
#
#                    if this_type == 'tbl' and caption:
#                        is_tabcap = True
#
#                    
#
#                    if last_type == 'img' and caption:
#                        path_print = path_print[:-1]
#                        path_print += " : " + caption.pop(0) + '\n' 
#                        has_caption = None
#
#                    doc_name += 1
#                    if not doc :
#                        doc = docx.Document('/data/projects/memory_compiler2/scripts/fire_stim_script/makeDoc/back.docx')
#                    if re.search('.docx', path_save):
#                        doc.element.body.remove(doc.element.body[0])
#                        doc.save(path_save)
#                        doc = docx.Document('/data/projects/memory_compiler2/scripts/fire_stim_script/makeDoc/back.docx')
#                    elif last_type == 'img':
#                        shutil.copy(path_of_image, path_save)
#
#
#                if this_type == 'img' :
#                    id_img =  re.sub(r'(.*<a:blip r:embed="|.*<v:imagedata r:id=")(.\w+)(.*)',r'\2',xml_index)
#                    for img in list_img:
#                        if img[0] == id_img :
#                            type_img = re.sub(r'(\w+).(.\w+)',r'\2', img[1])
#                            path_save = pathfolder + path_to_mkdir + '/' + str(doc_name) + '.' + type_img
#                            path_of_image = dirtrx + '/word/media/'+ img[1]
#                            last_type = this_type
#                            break
#                else:      
#                    try:
#                        if(xml_index != ''):
#                            output = xmlstring_to_oxmlelement(xml_index)
#                        path_save = pathfolder + path_to_mkdir + '/' + str(doc_name) + '.docx'
#
#                        doc.element.body[-2].addnext(output)
#                        last_type = 'p'
#                    except :
#                        print(xml_index+ '\n\n')
#                        exit()
#
#        if xml != '$' :
#            str_heading = re.findall(r'<w:t.*?>(.*?)</w:t>', xml[0])
#            str_heading = ''.join([str(elem) for elem in str_heading])
#            level_heading = re.sub(r'(.*w:val="Heading)(\d)(".*)', r'\2' ,xml.group())
#            last_xml_index = xml.span()[1]
#
#    if os.path.exists('pathname.txt') : 
#        os.remove('pathname.txt')
#    path_to_print = open(pathfile , 'a')
#    path_to_print.write(path_print)
#    print(path_print)
#    path_to_print.close()
#    shutil.rmtree(dirtrx)
#
#readdocx(woldfile)
