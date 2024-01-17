import os
import argparse
import re
import shutil
import zipfile
import random
import zipfile
import docx
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import WD_BREAK
from docx.shared import Inches,Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
###################################################################################################################

parser = argparse.ArgumentParser()
parser.add_argument('--tech', required=False, help="")
parser.add_argument('--type', required=False, help="")
cli_args = parser.parse_args()
# parser.add_argument('--tech', required=False, help="[ 80nm_90nm , 55nm_65nm , 12nm_16nm , 22nm_28nm , tsmc07nm ]")
# parser.add_argument('--type', required=False, help="[1PR , 2PR , DP , P2P , ROM , SPMB]")



def __main__ ():
    optionnote = open('/data/projects/memory_compiler2/REPOSITORY/Rampiler_doc/database/technote', 'r').readlines()
    option_tech = {}
    for line in optionnote : 
        print(line)
        line = line.strip()
        x = re.search(r'(\w+)\s?->\s?(\w+)', line).groups()
        option_tech[x[1]] = x[0]
    
    print(option_tech)
    
    # option_tech = {'tsmc90nm' : '80nm_90nm',
    #     'tsmc65nm' : '55nm_65nm',
    #     'tsmc16nm' : '12nm_16nm',
    #     'tsmc28nm' : '22nm_28nm',
    #     'tsmc07nm' : '06nm_07nm',
    #     'tsmc40nm' : '40nm',
    #     'tsmc03nm' : '03nm',
    #     'tsmc05nm' : '05nm'
    # }
    
    
    option_change = open('/data/projects/memory_compiler2/REPOSITORY/Rampiler_doc/database/list_change', 'r').readlines()
    option_change = [i.strip() for i in option_change]
    print(option_change)
    
    tech = get_input('tech')
    typ = get_input('type')
    if tech not in option_tech.keys():
        print('tech not exactly')
        exit()
        
    path_database = '/data/projects/memory_compiler2/REPOSITORY/Rampiler_doc/database/'
    path_total = path_database + 'tsmc_total/' + typ
    path_change = path_database + tech + '/' + typ
    
    if not os.path.exists(path_total):
        print('not exist path : %s'%path_total)
        exit()
  
    if not os.path.exists(path_change):
        print('not exist path : %s'%path_change)
        exit()
    
    
    # name  = get_input("name")
    lines_total = open(path_total+ '/tree', 'r', encoding='cp1252').readlines()
    lines_change = open(path_change + '/tree', 'r', encoding='cp1252').readlines()
    
    
    ######################################## change title, sub, date .... of output #################
    title = 'RAMpiler+ ®'
    subject = re.search(r'((\d+)nm_)?(\d+nm)', option_tech[tech]).groups()
    subject = (subject[2] if subject[0] == None else subject[1] + '/' + subject[2]) + ' ' +typ + ' Compiler'
    category = 'User Manual'
    revision = str(round(float(re.search(r'Rev_(\d+\.\d+)', lines_change[0]).groups()[0]) + 0.1, 2))
    revisionstr = 'Revision ' + revision
    print(revisionstr)
    date_update = datetime.datetime.today().strftime('%Y-%m-%d') + 'T00:00:00'
    
    output_name = 'DTI_%s_%s_RAMpilerplus_User_Manual_Rev_make'%(option_tech[tech],typ)
    shutil.copyfile('/data/projects/memory_compiler2/scripts/fire_stim_script/makeDoc/temp.docx', output_name + '.docx')
    
    
    with zipfile.ZipFile(output_name + '.docx','a') as propsfile:
        props = propsfile.read('docProps/core.xml').decode('utf-8')
        props = re.sub('\{title\}', title, props)
        props = re.sub('\{subject\}', subject, props)
        props = re.sub('\{category\}', category, props)
        props = re.sub('\{revision\}', revisionstr, props)
        propsfile.writestr('docProps/core.xml', props)
        
        dates = propsfile.read('customXml/item1.xml').decode('utf-8')
        dates = re.sub('1900-01-01T00:00:00', date_update, dates)
        propsfile.writestr('customXml/item1.xml', dates)
        
    doc = docx.Document(output_name + '.docx')
    
    num_of_figures = 1
    num_of_tables = 1

    lastlevel = -1
    start = False
    patharr = []
    data = []
    skip = False
    for i in range(2 , len(lines_total)):
        if not start:
            if re.search(r'TREE\s*',lines_total[i]):
                start = True
                continue
        if lines_total[i] == '\n' :
            continue
        
        line = lines_total[i]
        level = len(re.findall('\t|\s{4}', line)) + 1
        line = line.strip()
        
        names = re.split(r'\s+:\s+', line)
        folder_name = names[0]
        name_heading_total = names[1] if len(names) > 1 else ''
        name_heading_total_clean = re.sub(r'^\d+(\.(\d+))*\.(\s)?', '', name_heading_total) 
        if level > lastlevel :
            patharr.append(folder_name)
        else :
            patharr = patharr[:-(lastlevel - level + 1)]
            patharr.append(folder_name)
            
        lastlevel = level
        if skip:
            if level == level_skip:
                skip = False
            else:
                continue
            
        folder_total = '/'.join(patharr)
        path = path_total + '/' + folder_total
        if name_heading_total_clean in option_change:
            level_skip = level
            lastlevel = level
            skip = True
            start = False
            keep = False
            arr_change = []
            print(name_heading_total)
            for line in lines_change:
                if not start:
                    if re.search(r'TREE\s*',line):
                        start = True
                        continue
                if line == '\n' :
                    continue
                
                level = len(re.findall('\t|\s{4}', line)) + 1
                line = line.strip()
                names = re.split(r'\s+:\s+', line)
                
                folder_name = names[0]
                name_heading_change = names[1] if len(names) > 1 else ''
                name_heading_change_clean = re.sub(r'^\d+(\.(\d+))*\.(\s)?', '', name_heading_change)
                if level > lastlevel :
                    arr_change.append(folder_name)
                else :
                    arr_change = arr_change[:-(lastlevel - level + 1)]
                    arr_change.append(folder_name)
                folder_change = '/'.join(arr_change)
                path = path_change + '/' + folder_change
                if name_heading_change_clean == name_heading_total_clean :
                    keep = True
                if keep: 
                    if level == level_skip and name_heading_change_clean != name_heading_total_clean:
                        break
                    data.append([path,level,name_heading_change])
                
                lastlevel = level
        else:
            data.append([path,level,name_heading_total])
    lastlevel = -1
    doc.element.body.remove(doc.element.body[-2])
    last_heading = ''
    for db in data:
        pathname = db[0]
        level = db[1]
        name_heading = db[2]

        if level < lastlevel :
            run = doc.paragraphs[-1].add_run()
            run.add_break(WD_BREAK.PAGE)
        
        lastlevel = level

        print(pathname.ljust(95) ,'|', name_heading)

        isFolder = not re.search(r'\.(docx|png|jpeg|emf)', pathname)
        if isFolder :            
            if re.search(r'^\d+(\.(\d+))*', name_heading) :
                name_heading = re.sub(r'^\d+(\.(\d+))*\.(\s)?', '', name_heading)
                new_p = doc.add_heading(name_heading, level)
            else :
                new_p = doc.add_heading(name_heading, level)
                numPr = OxmlElement('w:numPR')
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'),'0')
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'),'0')
                numPr.append(ilvl)
                numPr.append(numId)
                new_p._element.pPr.append(numPr)
                if not re.search(r'REVISION', name_heading) :
                    new_p._element.pPr.jc_val = WD_ALIGN_PARAGRAPH.CENTER
        else: 
            if re.search(r'Figure\s+\\f[\.\:]', name_heading):
                if re.search(r'\.docx$',pathname):
                    doc = copy_docx(pathname, output_name, doc)
                else:
                    doc.add_picture(pathname, width=Inches(7.5))
                    paragraph = doc.paragraphs[-1]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                id_table_of_figures = random_id()
                figure_arr = re.findall(r'(Figure\s+)\\f(.*)', name_heading)[0]
                xml_string = '<w:p w:rsidR="' + id_table_of_figures +'" w:rsidRDefault="'+ id_table_of_figures +'" w:rsidP="'+ id_table_of_figures +'"><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t xml:space="preserve">'+ figure_arr[0] +'</w:t></w:r><w:fldSimple w:instr=" SEQ Figure \* ARABIC "><w:r><w:rPr><w:noProof/></w:rPr><w:t> ' + str(num_of_figures) + ' </w:t></w:r></w:fldSimple><w:r><w:t xml:space="preserve">'+ figure_arr[1] +'</w:t></w:r></w:p>'
                num_of_figures += 1 
                output = xmlstring_to_oxmlelement(xml_string)    
                doc.element.body[-2].addnext(output)

            if re.search(r'Table\s+\\t[\.\:]', name_heading):
                doc = copy_docx(pathname, output_name, doc)
                table_arr = re.findall(r'(Table\s+)\\t(.*)', name_heading)[0]
                id_table_of_tables = random_id()                    
                xml_string = '<w:p w:rsidR="'+ id_table_of_tables +'" w:rsidRDefault="'+ id_table_of_tables +'" w:rsidP="'+ id_table_of_tables +'"><w:pPr><w:pStyle w:val="Caption"/><w:keepNext/><w:jc w:val="center"/></w:pPr><w:proofErr w:type="gramStart"/><w:r><w:t xml:space="preserve">'+ table_arr[0] +' </w:t></w:r><w:fldSimple w:instr=" SEQ Table \* ARABIC "><w:r><w:rPr><w:noProof/></w:rPr><w:t> '+ str(num_of_tables) +' </w:t></w:r></w:fldSimple><w:proofErr w:type="gramEnd"/><w:r><w:t xml:space="preserve"> '+ table_arr[1] +'</w:t></w:r></w:p>'
                num_of_tables += 1
                output = xmlstring_to_oxmlelement(xml_string)    
                doc.element.body[-4].addnext(output)

            if name_heading == '':
                if re.search(r'\.(png|jpeg)$', pathname):
                    doc.add_picture(pathname, width=Inches(6.5))
                    paragraph = doc.paragraphs[-1]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc = copy_docx(pathname, output_name, doc)
                    
        # if last_heading == 'REVISION HISTORY':
        #     tablerow = [revision, datetime.datetime.today().strftime('%d/%m/%Y'), 'Merge multi doc file' , '' , 'Kientc0']
        #     doc = add_row_table(tablerow, doc)
        
        last_heading = name_heading
############################################################################################
    
    doc.save(output_name + '.docx')
    
    ################### đánh số lại toàn bộ id mũi tên và oval #########################
    with zipfile.ZipFile(output_name + '.docx', 'a') as doczip:
        data_xml_doc = doczip.read('word/document.xml').decode('utf-8')
        all_right_arrow = re.findall(r'id="(\d+)" name="((?!Picture).*?)"', data_xml_doc)
        new_id_arrow = 1000
        # all_right_arrow = list(all_right_arrow)
        # print(all_right_arrow)
        for id in all_right_arrow:
            if re.search( ' ' + id[0], id[1]):
                oldname = id[1]
                newname = id[1].replace(id[0], str(new_id_arrow))
                # print(oldname, newname) 
                data_xml_doc = re.sub(r'id="'+ id[0] +'" name="' + oldname, r'id="'+ str(new_id_arrow) +'" name="' + newname, data_xml_doc, count=1)
                data_xml_doc = re.sub(r' id="'+ oldname, r' id="'+ newname, data_xml_doc, count=1)
                new_id_arrow += 1
        
        doczip.writestr('word/document.xml', data_xml_doc)
    doc = docx.Document(output_name + '.docx')
    doc.save(output_name + '.docx')



def add_row_table(tablerow, doc) :
    table = doc.tables[-1]
    row = table.add_row().cells
    for i in range(5):
        row[i].text = tablerow[i]
        row[i].paragraphs[0].runs[0].font.size = Pt(10)
        row[i].paragraphs[0].paragraph_format.space_after = Inches(0)
        row[i].paragraphs[0].paragraph_format.space_before = Inches(0)
        if i < 2:
            row[i].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    return doc


def get_input(var_name):
    auto_input = getattr(cli_args, var_name, None)
    if auto_input :
        print("Auto input:", auto_input)
        return auto_input
    else:
        return input("Manual input: ")



def xmlstring_to_oxmlelement(string):
    queue = []
    save = []
    eval = ''
    first = ()
    array = re.finditer('<.+?>', string)
    index_last_text = 0
    for x in array:
        if x == '':
            continue
        else :
            text = x.group()
            if queue:         
                first = re.sub(r'.*:(.*)', r'\1', queue[0])    
                if queue[-1] == 'w:t':
                    if not re.search(r'<w:br/>',string[index_last_text:x.start()]):
                        eval += 't.text = \'\'\'' + string[index_last_text:x.start()] + '\'\'\'' + '\n'

                if re.search('/' + queue[-1], text) :
                    a = re.sub(r'.*:(.*)', r'\1', queue.pop())
                    while save:
                        eval += a + '.append(' + re.sub(r'.*:(.*)', r'\1', save.pop(0)) + ')' + '\n'
                    if queue:
                        eval += re.sub(r'.*:(.*)', r'\1', queue[-1]) + '.append(' + a + ')' + '\n'
                    continue
            
            text = re.sub('<|>', '', text)
            pins = text.split(' ',1)
            name_clr = re.sub('/', '', pins[0])
            name = re.sub(r'.*:(.*)', r'\1', name_clr)
            eval += name + ' = OxmlElement(\'' + name_clr + '\')'+'\n'
            queue.append(pins[0])

            if len(pins) > 1:
                matches = re.findall(r'\s?(.*?)=\"(.*?)\"', pins[1])
                for match in matches :
                    key = match[0]
                    value = match[1]
                    eval += name + '.set(qn(\'' + key + '\'),\'' + value + '\')' +'\n'
                
            if pins[0] == 'w:t':
                index_last_text = x.end()

            if re.search(r'\/', pins[-1]) :
                pin = re.sub('/', '' , pins[0])
                save.append(pin)
                queue.pop()
            
    eval = 'output = ()\n' + eval + 'output = ' + first
    loc = {}
    exec(eval, globals(), loc)
    output = loc['output']
    return output



def enter_click ():
    enter = OxmlElement('w:p')
    id = random_id()
    # <w:p w:rsidR="006948B6" w:rsidRDefault="006948B6"/>
    enter.set(qn('w:rsidR'), id)
    enter.set(qn('w:rsidRDefault'), id)
    return enter



def random_id ():
    return ''.join(random.choices('0123456789ABCDEF', k=8))


def list_id_pic(path):
    with zipfile.ZipFile(path, 'r') as zip:
        string_link = zip.read('word/_rels/document.xml.rels').decode('utf8')
    zip.close()
    list_img = re.findall(r'<Relationship Id="rId(\d+)".*?Target=".*?"\/>', string_link)
    list_img = list(map(int, list_img))
    list_img.sort()
    return list_img


def list_name_pic(path):
    with zipfile.ZipFile(path,'r') as zip:
        string_name = zip.namelist()
    zip.close()
    img_exist = []
    for name in string_name:
        img = re.search(r'word/media/image(\d+)\.',name)
        embedding = re.search(r'word/embeddings/.*?(\d+)\.',name)
        if img:
            img_exist.append(int(img.group(1)))
        if embedding:
            img_exist.append(int(embedding.group(1)))
    img_exist.sort()
    return img_exist

def copy_docx(path, path_save, doc) :
    doc_input = docx.Document(path)
    for doc_inside in doc_input.element.body:
        if str(type(doc_inside)) == '''<class 'docx.oxml.section.CT_SectPr'>''':
            break

        if re.search(r'="rId\d+"',doc_inside.xml):
            content_type_full = {
                'emf':'<Default Extension="emf" ContentType="image/x-emf"/>',
                'jpeg':'<Default Extension="jpeg" ContentType="image/jpeg"/>',
                "png" : '<Default Extension="png" ContentType="image/png"/>',
                "rels":'<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
                "vsd":'<Default Extension="vsd" ContentType="application/vnd.visio"/>',
                "vsdx":'<Default Extension="vsdx" ContentType="application/vnd.ms-visio.drawing"/>',
                "xml":'<Default Extension="xml" ContentType="application/xml"/>'
            }
            
            rid_xml_all_doc_before = re.findall(r'="rId(\d+)', doc.element.body.xml)
            rid_xml_all_before = {}
            for i in rid_xml_all_doc_before:
                if int(i) in rid_xml_all_before:
                    rid_xml_all_before[int(i)] += 1
                else :
                    rid_xml_all_before[int(i)] = 1
            doc.element.body[-2].addnext(doc_inside)
            doc.save(path_save + '.docx')

            with zipfile.ZipFile(path,'r') as inputzip:
                data_xml_rels_input = inputzip.read('word/_rels/document.xml.rels').decode('utf-8')
                
            with zipfile.ZipFile(path_save + '.docx','a') as doczip:
                data_xml_doc = doczip.read('word/document.xml').decode('utf-8')
                data_xml_rels_doc = doczip.read('word/_rels/document.xml.rels').decode('utf-8')
                content_type = doczip.read('[Content_Types].xml').decode('utf-8')
                rid_need_change = re.findall(r'="rId(\d+)"', doc_inside.xml)
                rid_need_change = list(dict.fromkeys(rid_need_change))
                # print(doc_inside.xml)
                # xml = re.sub(r'xmlns:(?!a)\w+=".*?"\s|\n\s+|\n','',doc_inside.xml)
                # xml = re.sub(r'\sxmlns:\w+=".*?">','>', xml)
                # print(xml)
                rid_xml_all_doc = re.findall(r'="rId(\d+)"', data_xml_doc)
                rid_xml_all = {}
                for i in rid_xml_all_doc:
                    if int(i) in rid_xml_all:
                        rid_xml_all[int(i)] += 1
                    else :
                        rid_xml_all[int(i)] = 1

                num_id_img_exist = list_id_pic(path_save + '.docx')
                num_id_path = list_name_pic(path_save + '.docx')
                

                if num_id_img_exist and num_id_path:
                    num_last_id = num_id_img_exist[-1]
                    num_last_id_path = num_id_path[-1]
                else:
                    num_last_id = 0 
                    num_last_id_path = 0
                # print(rid_need_change)
                # print(rid_xml_all_before)
                for rid in rid_need_change :
                    num_last_id += 1
                    num_last_id_path += 1
                    
                    if 21 <= int(rid) <= 22:
                        data_xml_doc = re.sub(r'="rId'+ rid + '"', r'="rId'+ str(num_last_id) + '"', data_xml_doc)
                        
                    if int(rid) in list(rid_xml_all_before.keys()):
                        if 17 <= int(rid) <= 20 :
                            data_xml_doc = re.sub(r'="rId'+ rid + '"', r'="rId'+ str(num_last_id) + '"', data_xml_doc, count= len(re.findall(r'="rId'+ rid, data_xml_doc)) - 1)
                        else:
                            data_xml_doc = re.sub(r'="rId'+ rid + '"', r'="rId'+ str(num_last_id) + '"', data_xml_doc)
                            data_xml_doc = re.sub(r'="rId'+ str(num_last_id) + '"', r'="rId' + rid + '"', data_xml_doc, count=rid_xml_all_before[int(rid)])
                    data_xml_rels = re.findall(r'<Relationship Id="rId'+ rid + '"' + r'.*?Target=".*?"/>',  data_xml_rels_input)
                    
                    # if len(data_xml_rels) == 0:
                    #     continue
                    data_xml_rels = data_xml_rels[0]

                    data_xml_rels_to_doc = re.sub(r'<Relationship Id="rId'+ rid + r'"(.*?)Target="(.*?)\d+(\..*?"/>)' , r'<Relationship Id="rId'+ str(num_last_id)+ r'"\1Target="\g<2>' + str(num_last_id_path) + r'\3',  data_xml_rels)
                    data_xml_rels_doc = re.sub(r'</Relationships>', data_xml_rels_to_doc + '</Relationships>', data_xml_rels_doc)
                    
                    path_of_img = re.findall(r'<Relationship Id="rId'+ rid + r'.*?Target="(.*?)"/>',  data_xml_rels_input)[0]
                    new_path_img = re.sub(r'\d+\.', str(num_last_id_path) + '.', path_of_img)
                        
                    type_content = re.sub(r'(.*)\.(\w+)',r'\2', path_of_img)
                    
                    if not re.search(content_type_full[type_content], content_type):
                        content_type = re.sub(r'<Types(.*?)>',r'<Types\1>'+ content_type_full[type_content] ,content_type)
                        doczip.writestr('[Content_Types].xml', content_type)
                    
                    doczip.writestr('word/' + new_path_img, zipfile.ZipFile(path, 'r').open('word/' + path_of_img).read())
                doczip.writestr('word/document.xml', data_xml_doc)
                doczip.writestr('word/_rels/document.xml.rels', data_xml_rels_doc)
                

            doc = docx.Document(path_save + '.docx')
        else :
            doc.element.body[-2].addnext(doc_inside)

        if str(type(doc_inside)) == '''<class 'docx.oxml.table.CT_Tbl'>''':
            doc.element.body[-2].addnext(enter_click())

    return doc


__main__()





#def copy_docx(path) :
#    string = ()
#    with zipfile.ZipFile(path, 'r') as zip:
#        string = zip.read('word/document.xml').decode('utf8')
#    string = str(string)
#    xml_string = re.findall('<w:tbl>.+?</w:tbl>|<w:p.+?</w:p>' , string)
#    for xml in xml_string :
#        output = xmlstring_to_oxmlelement(xml)    
#        if re.search(r'<w:p.+?</w:p>',xml):
#            if re.search(r'<w:tbl>', doc.element.body[-3].xml) :
#                doc.element.body.remove(doc.element.body[-2])
#
#        doc.element.body[-2].addnext(output)
#        if re.search(r'<w:tbl>',xml) : 
#            doc.element.body[-2].addnext(enter_click())
#
#            
#
#def add_text_to_docx(path) :
#    copy_docx(path)
#
#def add_picture_to_docx(path , name_picture) :
#    global num_of_figures
#    id_table_of_figures = random_id()
#    doc.add_picture(path, width=Inches(7.5))
#    paragraph = doc.paragraphs[-1]
#    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    paragraph._p.set(qn('w:rsidRDefault'),id_table_of_figures)
#    if name_picture != '':
#        xml_string = '<w:p w:rsidR="' + id_table_of_figures +'" w:rsidRDefault="'+ id_table_of_figures +'" w:rsidP="'+ id_table_of_figures +'"><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t xml:space="preserve">Figure </w:t></w:r><w:fldSimple w:instr=" SEQ Figure \* ARABIC "><w:r><w:rPr><w:noProof/></w:rPr><w:t> ' + str(num_of_figures) + ' </w:t></w:r></w:fldSimple><w:r><w:t>: '+ name_picture +'</w:t></w:r></w:p>'
#        num_of_figures += 1 
#        output = xmlstring_to_oxmlelement(xml_string)    
#        doc.paragraphs[-1]._p.addnext(output)
#
#
#def add_table_to_docx(path, name_table) :
#    global num_of_tables
#    id_table_of_tables = random_id()
#    copy_docx(path)
#    xml_string = '<w:p w:rsidR="'+ id_table_of_tables +'" w:rsidRDefault="'+ id_table_of_tables +'" w:rsidP="'+ id_table_of_tables +'"><w:pPr><w:pStyle w:val="Caption"/><w:keepNext/><w:jc w:val="center"/></w:pPr><w:proofErr w:type="gramStart"/><w:r><w:t xml:space="preserve">Table </w:t></w:r><w:fldSimple w:instr=" SEQ Table \* ARABIC "><w:r><w:rPr><w:noProof/></w:rPr><w:t> '+ str(num_of_tables) +' </w:t></w:r></w:fldSimple><w:r><w:t>.</w:t></w:r><w:proofErr w:type="gramEnd"/><w:r><w:t xml:space="preserve"> '+ name_table +'</w:t></w:r></w:p>'
#    num_of_tables += 1
#    output = xmlstring_to_oxmlelement(xml_string)    
#    doc.paragraphs[-2]._p.addnext(output)


#def get_input(var_name):
#    auto_input = getattr(cli_args, var_name, None)
#    if auto_input :
#        print("Auto input:", auto_input)
#        return auto_input
#    else:
#        return input("Manual input: ")
####################################################################################################################
#
#parser = argparse.ArgumentParser()
#parser.add_argument('--path', default= 'pathname.txt' , required=False)
## parser.add_argument('--name', default=None, required=False)
#cli_args = parser.parse_args()
#
#
#path = get_input("path")
## name  = get_input("name")
#
#path = open(path, 'r', encoding='cp1252')
#lines = path.readlines()
#path.close()
#pathname = re.sub('^__path__ = |\n', '' , lines[0])
#output_name = lines[1]
#output_name = re.sub(r'^name_of_output = |\n','',output_name)
#
#shutil.copyfile('/data/projects/memory_compiler2/scripts/fire_stim_script/makeDoc/temp.docx', output_name + '.docx')
#doc = docx.Document(output_name + '.docx')
#
#num_of_figures = 1
#num_of_tables = 1
#
#lastlevel = -1
#folder = pathname
#tree_line_start = 0
#for n in range(0 , len(lines)):
#    if re.search(r'^TREE\s*:', lines[n]):
#        tree_line_start = n + 1
#        break
#
#for i in range(tree_line_start , len(lines)):
#    if lines[i] == '\n' :
#        continue
#    else :
#        level = len(re.findall('\t|\s{4}', lines[i])) + 1
#        if level <= lastlevel : 
#            pathname = re.sub(r'(.*)(/.*){'+ str(lastlevel - level + 1) + '}', r'\1', pathname)
#
#        if level < lastlevel or level == 1:
#            run = doc.paragraphs[-1].add_run()
#            run.add_break(WD_BREAK.PAGE)
#        
#        name = re.sub('\n|\t|\s{4}', '' , lines[i])
#        names = name.split(' : ')
#        folder_name = names[0]
#        name_header = ''
#        if len(names) > 1 :
#            name_header = names[1]
#        pathname = pathname + '/' + folder_name
#        if (re.search(r'\.emf', pathname)) :
#          print("ERROR: Wrong format {} -> <*.png|*.jpg>".format(pathname))
#
#        isFolder = not re.search('.docx|.png|.jpg|.emf', folder_name)
#
#        if isFolder :
#            if re.search(r'^\d+(\.(\d+))*', name_header) :
#                name_header = re.sub(r'^\d+(\.(\d+))*\.(\s)?', '', name_header)
#                new_p = doc.add_heading(name_header, level)
#            else :
#                new_p = doc.add_heading(name_header, level)
#                numPr = OxmlElement('w:numPR')
#                ilvl = OxmlElement('w:ilvl')
#                ilvl.set(qn('w:val'),'0')
#                numId = OxmlElement('w:numId')
#                numId.set(qn('w:val'),'0')
#                numPr.append(ilvl)
#                numPr.append(numId)
#                new_p._element.pPr.append(numPr)
##                new_p._element.pPr.jc_val = WD_ALIGN_PARAGRAPH.CENTER
#
#        elif re.search('.docx', folder_name) and name_header != '':
#
#            add_table_to_docx(pathname , name_header)
#
#        elif re.search('.docx',folder_name):
#
#            add_text_to_docx(pathname)
#
#        elif re.search('.jpg|.png', folder_name) :
#            
#            add_picture_to_docx(pathname, name_header)
#        
#        lastlevel = level
#
#doc.save(output_name + '.docx')
#print("Info OutPut: " + output_name + '.docx')
#
